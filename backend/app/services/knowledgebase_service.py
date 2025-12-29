from __future__ import annotations

from typing import Dict, List, Optional, Tuple
import os
import io
import math
import mimetypes

from openai import AsyncOpenAI

from app.dependencies import _get_openai_client
from app.repository import knowledgebase_repository as repo


EMBEDDING_MODEL = "text-embedding-3-large"


def _split_text(text: str, *, chunk_size: int = 1200, chunk_overlap: int = 200) -> List[str]:
    if chunk_size <= 0:
        return [text]
    chunks: List[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunks.append(text[start:end])
        if end == length:
            break
        start = end - chunk_overlap
        if start < 0:
            start = 0
    return chunks


def _cosine_similarity(vec_a: List[float], vec_b: List[float]) -> float:
    if not vec_a or not vec_b:
        return 0.0
    if len(vec_a) != len(vec_b):
        # Best-effort; shouldn't happen if model is consistent
        n = min(len(vec_a), len(vec_b))
        vec_a = vec_a[:n]
        vec_b = vec_b[:n]
    dot = sum(a * b for a, b in zip(vec_a, vec_b))
    norm_a = math.sqrt(sum(a * a for a in vec_a))
    norm_b = math.sqrt(sum(b * b for b in vec_b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


async def _embed_texts(client: AsyncOpenAI, texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
    resp = await client.embeddings.create(model=EMBEDDING_MODEL, input=texts)
    return [item.embedding for item in resp.data]  # type: ignore[attr-defined]


async def process_uploaded_file(
    *,
    storage_path: str,
    original_name: str,
    mime_type: str,
    description: Optional[str],
    is_global: bool,
    file_hash: Optional[str] = None,
) -> str:
    # Read file content and attempt to infer a document title
    title, text = await _extract_title_and_text(storage_path, mime_type)
    client = _get_openai_client()

    # Optionally generate a short description from content when not provided
    auto_description: Optional[str] = None
    if not description:
        try:
            auto_description = await _generate_description(client, text)
        except Exception:
            auto_description = None

    # Create document entry first
    stat = os.stat(storage_path)
    filename = os.path.basename(storage_path)
    doc_id = await repo.create_document(
        filename=filename,
        original_name=original_name,
        size=stat.st_size,
        mime_type=mime_type,
        storage_path=storage_path,
        description=description or auto_description,
        is_global=is_global,
        title=title,
        file_hash=file_hash,
    )

    # Compute a document-level embedding from title/description or first chunk
    title_desc = f"{(title or original_name)}\n\n{(description or auto_description or '')}".strip()
    if title_desc:
        doc_embeds = await _embed_texts(client, [title_desc])
        if doc_embeds:
            await repo.update_document_embedding(
                document_id=doc_id,
                embedding=doc_embeds[0],
                model=EMBEDDING_MODEL,
                dimensions=len(doc_embeds[0]),
            )

    # Chunk text and embed
    chunks = _split_text(text)
    embeddings = await _embed_texts(client, chunks)
    chunk_payload: List[Tuple[int, str, List[float]]] = [
        (idx, chunk, embeddings[idx]) for idx, chunk in enumerate(chunks)
    ]
    await repo.insert_chunk_embeddings(
        document_id=doc_id,
        chunks=chunk_payload,
        model=EMBEDDING_MODEL,
        dimensions=len(embeddings[0]) if embeddings else 0,
    )

    return doc_id


async def _generate_description(client: AsyncOpenAI, text: str) -> str:
    if not text:
        return ""
    sample = text[:6000]
    prompt = (
        "You will be given the beginning of a document. In 1-2 concise sentences, "
        "write a neutral summary describing the document's purpose and content for a file library UI. "
        "Avoid conjecture; do not exceed 40 words."
    )
    resp = await client.responses.create(
        model="gpt-4o-mini",
        input=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": sample},
        ],
    )
    try:
        # Responses API: first output_text helper
        out = resp.output_text  # type: ignore[attr-defined]
        return out.strip() if out else ""
    except Exception:
        # Fallback: try content array
        try:
            text_parts = []
            for out in getattr(resp, "output", []) or []:  # type: ignore[attr-defined]
                val = getattr(out, "content", None)
                if isinstance(val, str):
                    text_parts.append(val)
            return (" ".join(text_parts)).strip()[:300]
        except Exception:
            return ""


async def _extract_text(path: str, mime_type: str) -> str:
    # Handle simple text types
    if mime_type.startswith("text/") or mime_type in ("application/json",):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    # Basic PDF support via pypdf if installed
    if mime_type == "application/pdf":
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(path)
            text = []
            for page in reader.pages:
                try:
                    text.append(page.extract_text() or "")
                except Exception:
                    continue
            return "\n".join(text)
        except Exception:
            return ""

    # Basic DOCX support via python-docx if installed
    if mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
        try:
            import docx  # type: ignore

            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""

    # Fallback: try to guess and read as text
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


async def _extract_title_and_text(path: str, mime_type: str) -> Tuple[Optional[str], str]:
    """Best-effort extraction of a human-friendly title and full text.

    - For PDFs, prefer metadata title; otherwise use the first non-empty line on page 1.
    - For DOCX, use the first non-empty paragraph as title.
    - For plain text, use the first non-empty line.
    - Fallback to None when we cannot derive a title.
    """
    title: Optional[str] = None

    # PDF path (try metadata and first page headers)
    if mime_type == "application/pdf":
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(path)
            # Extract text across all pages
            pages_text: List[str] = []
            first_page_text: str = ""
            for idx, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    pages_text.append(page_text)
                    if idx == 0:
                        first_page_text = page_text
                except Exception:
                    continue

            # Title from metadata if present
            try:
                meta = getattr(reader, "metadata", None)
                meta_title = None
                if meta is not None:
                    meta_title = getattr(meta, "title", None)
                if isinstance(meta_title, str):
                    cleaned = meta_title.strip()
                    title = cleaned if cleaned else None
            except Exception:
                title = None

            # Fallback: first non-empty reasonably short line from first page
            if not title and first_page_text:
                for line in (first_page_text.splitlines() or []):
                    candidate = (line or "").strip()
                    if 3 <= len(candidate) <= 140:
                        title = candidate
                        break

            return title, "\n".join(pages_text)
        except Exception:
            # Fall through to generic extraction
            pass

    # DOCX path
    if mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
        try:
            import docx  # type: ignore

            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            text = "\n".join(paragraphs)
            for p in paragraphs:
                candidate = (p or "").strip()
                if candidate:
                    title = candidate[:140]
                    break
            return title, text
        except Exception:
            pass

    # Plain text and everything else: best-effort
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        if content:
            for line in content.splitlines():
                candidate = (line or "").strip()
                if candidate:
                    title = candidate[:140]
                    break
        return title, content
    except Exception:
        return None, ""


async def reextract_title_for_document(*, storage_path: str, mime_type: str) -> Optional[str]:
    """Re-extract a best-effort title from the stored document without re-embedding.

    Returns the new title if found, otherwise None.
    """
    title, _ = await _extract_title_and_text(storage_path, mime_type)
    return title


async def retrieve_context(
    *,
    query: str,
    mode: str,  # 'all' or 'file'
    file_id: Optional[str] = None,
    max_docs: int = 5,
    max_chunks: int = 8,
) -> Dict[str, any]:
    client = _get_openai_client()
    # Embed query
    query_vec = (await _embed_texts(client, [query]))[0]

    # Determine candidate documents
    candidate_doc_ids: List[str] = []
    if mode == "file" and file_id:
        candidate_doc_ids = [file_id]
    else:
        # Rank documents by similarity of doc_embedding
        docs = await repo.list_all_doc_embeddings()
        scored: List[Tuple[str, float, str]] = []  # (doc_id, score, title)
        for d in docs:
            emb = d.get("doc_embedding") or []
            sim = _cosine_similarity(query_vec, emb)
            scored.append((d["id"], sim, d.get("filename", "")))
        scored.sort(key=lambda x: x[1], reverse=True)
        candidate_doc_ids = [d[0] for d in scored[:max_docs] if d[1] > 0.2]

    if not candidate_doc_ids:
        return {"chunks": [], "doc_ids": []}

    # Pull chunks for those documents and rank by similarity
    chunks = await repo.list_chunks_for_documents(candidate_doc_ids)
    scored_chunks: List[Tuple[float, Dict[str, any]]] = []
    for ch in chunks:
        emb = ch.get("embedding") or []
        sim = _cosine_similarity(query_vec, emb)
        scored_chunks.append((sim, ch))
    scored_chunks.sort(key=lambda x: x[0], reverse=True)
    top = [chunk for sim, chunk in scored_chunks[:max_chunks] if sim > 0.2]
    return {"chunks": top, "doc_ids": candidate_doc_ids}


def build_rag_messages(original_messages: List[Dict[str, str]], *, context_chunks: List[Dict[str, any]]) -> List[Dict[str, str]]:
    if not context_chunks:
        return original_messages
    # Build a single system message containing the retrieved context
    ctx_lines = []
    for ch in context_chunks:
        ref = f"doc:{ch['document_id']}#chunk:{ch['chunk_index']}"
        snippet = ch["text"]
        ctx_lines.append(f"[{ref}]\n{snippet}")
    context_block = "\n\n".join(ctx_lines)
    system_msg = {
        "role": "system",
        "content": (
            "You are a precise assistant. Use the provided knowledgebase context when relevant. "
            "If the context is insufficient, say so and answer from your general knowledge.\n\n"
            f"Context:\n{context_block}"
        ),
    }
    return [system_msg] + original_messages


